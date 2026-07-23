from pathlib import Path

from docx import Document

from backend.app.services.docx_extractor import extract_docx_text


def test_extract_docx_text_strips_empty_paragraphs(tmp_path: Path) -> None:
    docx_path = tmp_path / "sample.docx"
    document = Document()
    document.add_paragraph("First paragraph")
    document.add_paragraph("   ")
    document.add_paragraph("Second paragraph")
    document.save(docx_path)

    extracted = extract_docx_text(docx_path)

    assert extracted.text == "First paragraph\nSecond paragraph"
    assert extracted.char_count == len(extracted.text)
    assert extracted.paragraphs == ["First paragraph", "Second paragraph"]
    assert extracted.headline == "First paragraph"
    assert extracted.subheadline is None
    assert extracted.body_text == "Second paragraph"
    assert extracted.headline_status == "inferred"


def test_extract_docx_text_infers_headline_subheadline_and_body(
    tmp_path: Path,
) -> None:
    docx_path = tmp_path / "sample.docx"
    document = Document()
    document.add_paragraph("Headline")
    document.add_paragraph("Subheadline")
    document.add_paragraph("Body paragraph one")
    document.add_paragraph("Body paragraph two")
    document.save(docx_path)

    extracted = extract_docx_text(docx_path)

    assert extracted.headline == "Headline"
    assert extracted.subheadline == "Subheadline"
    assert extracted.body_text == "Body paragraph one\nBody paragraph two"
    assert extracted.word_count == 8
    assert extracted.headline_status == "inferred"
    assert extracted.subheadline_status == "inferred"


def test_extract_docx_text_does_not_treat_long_lede_as_headline(
    tmp_path: Path,
) -> None:
    docx_path = tmp_path / "sample.docx"
    lede = " ".join(f"word{i}" for i in range(30)) + "."
    document = Document()
    document.add_paragraph(lede)
    document.add_paragraph("Second body paragraph")
    document.save(docx_path)

    extracted = extract_docx_text(docx_path)

    assert extracted.headline is None
    assert extracted.headline_candidate == lede
    assert extracted.headline_status == "not_present"
    assert extracted.body_text == f"{lede}\nSecond body paragraph"
